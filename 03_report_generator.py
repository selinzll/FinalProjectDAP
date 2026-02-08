import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Circle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random

def create_radar_chart(company_data, dimensions, filename):
    before_scores = [company_data[f'DimScore_{dim}_Before'] for dim in dimensions]
    after_scores = [company_data[f'DimScore_{dim}_After'] for dim in dimensions]
    
    categories = dimensions
    N = len(categories)
    
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    before_scores += before_scores[:1]
    after_scores += after_scores[:1]
    angles += angles[:1]
    
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))
    
    ax.plot(angles, before_scores, 'o-', linewidth=2, label='Before', color='#E74C3C', markersize=8)
    ax.fill(angles, before_scores, alpha=0.15, color='#E74C3C')
    
    ax.plot(angles, after_scores, 'o-', linewidth=2, label='After', color='#27AE60', markersize=8)
    ax.fill(angles, after_scores, alpha=0.15, color='#27AE60')
    
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, size=11, weight='bold')
    ax.set_ylim(0, 100)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_yticklabels(['20', '40', '60', '80', '100'], size=9, color='gray')
    ax.grid(True, linestyle='--', alpha=0.5)
    ax.set_facecolor('#F8F9FA')
    
    plt.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=11, frameon=True)
    plt.title(f"{company_data['Company_Name_Before']}\nDigital Maturity Assessment", 
              size=14, weight='bold', pad=20)
    
    plt.tight_layout()
    plt.savefig(filename, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    
    return filename

def generate_pdf_report(company_data, dimensions, output_path):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    header = doc.add_heading('Digital Maturity Assessment Report', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.runs[0]
    header_run.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_paragraph()
    
    company_info = doc.add_paragraph()
    company_info.add_run(f"Company: ").bold = True
    company_info.add_run(f"{company_data['Company_Name_Before']}\n")
    company_info.add_run(f"Sector: ").bold = True
    company_info.add_run(f"{company_data['Sector_Before']}\n")
    company_info.add_run(f"Size: ").bold = True
    company_info.add_run(f"{company_data['Size_Before']}\n")
    company_info.add_run(f"Country: ").bold = True
    company_info.add_run(f"{company_data['Country_Before']}\n")
    company_info.add_run(f"Assessment Date: ").bold = True
    company_info.add_run(f"{company_data['Assessment_Date_After']}\n")
    
    doc.add_paragraph()
    
    doc.add_heading('Executive Summary', 1)
    
    summary = doc.add_paragraph()
    
    overall_before = company_data['Overall_Maturity_Before']
    overall_after = company_data['Overall_Maturity_After']
    growth = overall_after - overall_before
    pct_growth = (growth / overall_before * 100) if overall_before > 0 else 0
    level = company_data['Maturity_Level_After']
    
    summary.add_run(
        f"This report presents the digital maturity assessment results for "
        f"{company_data['Company_Name_Before']}. "
    )
    
    if growth > 0:
        summary.add_run(
            f"The organization has achieved significant progress, improving from a baseline "
            f"maturity score of {overall_before:.1f} to {overall_after:.1f}, representing a "
        )
        growth_run = summary.add_run(f"+{growth:.1f} point ({pct_growth:.1f}%) increase")
        growth_run.bold = True
        growth_run.font.color.rgb = RGBColor(39, 174, 96)
        summary.add_run(f". ")
    else:
        summary.add_run(
            f"The organization shows a maturity score of {overall_after:.1f}, compared to "
            f"a baseline of {overall_before:.1f}. "
        )
    
    summary.add_run(f"This performance classifies the organization as a ")
    level_run = summary.add_run(f"{level}")
    level_run.bold = True
    level_run.font.color.rgb = RGBColor(52, 73, 94)
    summary.add_run(f" in digital maturity.")
    
    doc.add_paragraph()
    
    doc.add_heading('Key Performance Indicators', 1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Before'
    hdr_cells[2].text = 'After'
    hdr_cells[3].text = 'Change'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    row = table.add_row().cells
    row[0].text = 'Overall Maturity Score'
    row[1].text = f'{overall_before:.1f}'
    row[2].text = f'{overall_after:.1f}'
    row[3].text = f'+{growth:.1f}' if growth >= 0 else f'{growth:.1f}'
    
    for dim in dimensions:
        row = table.add_row().cells
        row[0].text = dim
        before = company_data[f'DimScore_{dim}_Before']
        after = company_data[f'DimScore_{dim}_After']
        delta = after - before
        row[1].text = f'{before:.1f}'
        row[2].text = f'{after:.1f}'
        row[3].text = f'+{delta:.1f}' if delta >= 0 else f'{delta:.1f}'
    
    doc.add_paragraph()
    
    doc.add_heading('Visual Performance Profile', 1)
    
    doc.add_paragraph(
        "The radar chart below illustrates your organization's digital maturity across "
        "all six assessment dimensions, comparing baseline (Before) and current (After) states."
    )
    
    chart_path = f'/Users/zseli/PycharmProjects/FinalProjectDAP/radar_{company_data["Company_ID"]}.png'
    create_radar_chart(company_data, dimensions, chart_path)
    
    doc.add_picture(chart_path, width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading('Dimensional Analysis', 1)
    
    dim_deltas = [(dim, company_data[f'{dim}_Delta']) for dim in dimensions]
    dim_deltas.sort(key=lambda x: x[1], reverse=True)
    
    strongest_dim, strongest_growth = dim_deltas[0]
    weakest_dim, weakest_growth = dim_deltas[-1]
    
    analysis_para = doc.add_paragraph()
    analysis_para.add_run("Strengths: ").bold = True
    analysis_para.add_run(
        f"The {strongest_dim} dimension shows the highest improvement (+{strongest_growth:.1f} points), "
        f"indicating strong progress in this area. "
    )
    
    if weakest_growth < 5:
        analysis_para.add_run("\n\nAreas for Focus: ").bold = True
        analysis_para.add_run(
            f"The {weakest_dim} dimension requires attention, with limited growth "
            f"({weakest_growth:+.1f} points). Targeted interventions in this area could yield significant benefits."
        )
    
    doc.add_paragraph()
    
    doc.add_heading('Strategic Recommendations', 1)
    
    recommendations = []
    
    if weakest_growth < 10:
        recommendations.append(
            f"Develop a focused improvement plan for {weakest_dim}, including skills training, "
            f"technology investments, and process re-engineering."
        )
    
    if overall_after < 45:
        recommendations.append(
            "Establish foundational digital capabilities through basic technology adoption "
            "and staff training programs."
        )
    elif overall_after < 75:
        recommendations.append(
            "Build on current capabilities by implementing advanced technologies and "
            "data-driven decision making processes."
        )
    else:
        recommendations.append(
            "Maintain leadership position through continuous innovation and adoption of "
            "emerging technologies."
        )
    
    recommendations.append(
        "Conduct quarterly maturity assessments to track progress and adjust strategies accordingly."
    )
    
    for i, rec in enumerate(recommendations, 1):
        para = doc.add_paragraph(rec, style='List Number')
    
    doc.add_paragraph()
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "This report is based on the EU Open Digital Maturity Assessment Tool (DMAT) Framework\n"
        "Generated by Digital Transformation Analytics Suite"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(output_path)
    print(f"✓ Generated report: {output_path}")
    
    return output_path

def generate_sample_reports(n_samples=5):
    print("=" * 70)
    print("GENERATING SAMPLE PDF REPORTS")
    print("=" * 70)
    
    df_before = pd.read_excel('/Users/zseli/PycharmProjects/FinalProjectDAP/rawdma_before.xlsx')
    df_after = pd.read_excel('/Users/zseli/PycharmProjects/FinalProjectDAP/rawdma_after.xlsx')
    
    df = df_before.merge(df_after, on='Company_ID', suffixes=('_Before', '_After'))
    
    dimensions = ['Strategy', 'Readiness', 'HumanCentric', 'DataMgmt', 'AutomationAI', 'GreenDigital']
    
    for dim in dimensions:
        df[f'{dim}_Delta'] = df[f'DimScore_{dim}_After'] - df[f'DimScore_{dim}_Before']
    
    df['Overall_Delta'] = df['Overall_Maturity_After'] - df['Overall_Maturity_Before']
    
    top_performers = df.nlargest(2, 'Overall_Delta')
    bottom_performers = df.nsmallest(2, 'Overall_Delta')
    random_sample = df.sample(n=max(0, n_samples-4), random_state=42)
    
    sample_df = pd.concat([top_performers, bottom_performers, random_sample])
    
    print(f"\nGenerating {len(sample_df)} reports...")
    
    generated_reports = []
    
    for idx, row in sample_df.iterrows():
        company_id = row['Company_ID']
        company_name = row['Company_Name_Before']
        
        output_path = f'/Users/zseli/PycharmProjects/FinalProjectDAP/Report_{company_id}_{company_name.replace(" ", "_")}.docx'
        
        try:
            generate_pdf_report(row, dimensions, output_path)
            generated_reports.append(output_path)
        except Exception as e:
            print(f"✗ Error generating report for {company_name}: {e}")
    
    print("\n" + "=" * 70)
    print(f"REPORT GENERATION COMPLETE: {len(generated_reports)} reports created")
    print("=" * 70)
    
    print("\nGenerated Reports:")
    print("-" * 70)
    for report_path in generated_reports:
        filename = report_path.split('/')[-1]
        print(f"  {filename}")
    
    return generated_reports

if __name__ == "__main__":
    reports = generate_sample_reports(n_samples=5)